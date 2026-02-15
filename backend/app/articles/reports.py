"""
Report generation and export for articles and content.
Supports PDF and Word export with executive/technical summaries.
"""
from fastapi import APIRouter, Depends, HTTPException, Response
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.auth.dependencies import get_current_user
from app.models import User, FetchedContent
from typing import List, Optional
from datetime import datetime
from io import BytesIO
from app.core.logging import logger

router = APIRouter(prefix="/articles/reports", tags=["article-reports"])


# ============================================================================
# Report Generation Functions
# ============================================================================

def generate_pdf_report(
    title: str,
    content: str,
    executive_summary: Optional[str],
    technical_summary: Optional[str],
    iocs: Optional[dict],
    metadata: Optional[dict]
) -> bytes:
    """Generate PDF report using ReportLab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib import colors

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()

        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1890ff'),
            spaceAfter=30
        )
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2 * inch))

        # Executive Summary
        if executive_summary:
            story.append(Paragraph("<b>Executive Summary</b>", styles['Heading2']))
            story.append(Paragraph(executive_summary, styles['BodyText']))
            story.append(Spacer(1, 0.3 * inch))

        # Technical Summary
        if technical_summary:
            story.append(Paragraph("<b>Technical Summary</b>", styles['Heading2']))
            story.append(Paragraph(technical_summary, styles['BodyText']))
            story.append(Spacer(1, 0.3 * inch))

        # IOCs
        if iocs:
            story.append(Paragraph("<b>Indicators of Compromise (IOCs)</b>", styles['Heading2']))

            for ioc_type, values in iocs.items():
                if values and isinstance(values, list) and len(values) > 0:
                    story.append(Paragraph(f"<b>{ioc_type.upper()}:</b>", styles['Heading3']))
                    for value in values[:20]:  # Limit to 20 per type
                        story.append(Paragraph(f"â€¢ {value}", styles['BodyText']))
                    story.append(Spacer(1, 0.1 * inch))

        # Metadata
        story.append(PageBreak())
        story.append(Paragraph("<b>Document Information</b>", styles['Heading2']))

        if metadata:
            meta_data = [
                ['Field', 'Value'],
                ['Format', metadata.get('content_format', 'Unknown')],
                ['Generated', datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')],
                ['Source URL', metadata.get('url', 'N/A')[:60]]
            ]

            meta_table = Table(meta_data, colWidths=[2 * inch, 4 * inch])
            meta_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(meta_table)

        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    except ImportError:
        raise Exception("reportlab not installed. Install with: pip install reportlab")
    except Exception as e:
        raise Exception(f"PDF generation failed: {str(e)}")


def generate_word_report(
    title: str,
    content: str,
    executive_summary: Optional[str],
    technical_summary: Optional[str],
    iocs: Optional[dict],
    metadata: Optional[dict]
) -> bytes:
    """Generate Word (.docx) report using python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Executive Summary
        if executive_summary:
            doc.add_heading('Executive Summary', 1)
            doc.add_paragraph(executive_summary)

        # Technical Summary
        if technical_summary:
            doc.add_heading('Technical Summary', 1)
            doc.add_paragraph(technical_summary)

        # IOCs
        if iocs:
            doc.add_heading('Indicators of Compromise (IOCs)', 1)

            for ioc_type, values in iocs.items():
                if values and isinstance(values, list) and len(values) > 0:
                    doc.add_heading(ioc_type.upper(), 2)
                    for value in values[:20]:
                        doc.add_paragraph(value, style='List Bullet')

        # Metadata
        doc.add_page_break()
        doc.add_heading('Document Information', 1)

        if metadata:
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Light Grid Accent 1'

            cells = table.rows[0].cells
            cells[0].text = 'Field'
            cells[1].text = 'Value'

            cells = table.rows[1].cells
            cells[0].text = 'Format'
            cells[1].text = metadata.get('content_format', 'Unknown')

            cells = table.rows[2].cells
            cells[0].text = 'Generated'
            cells[1].text = datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')

            cells = table.rows[3].cells
            cells[0].text = 'Source URL'
            cells[1].text = metadata.get('url', 'N/A')[:60]

        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except ImportError:
        raise Exception("python-docx not installed. Install with: pip install python-docx")
    except Exception as e:
        raise Exception(f"Word generation failed: {str(e)}")


# ============================================================================
# Routes
# ============================================================================

@router.get("/{content_id}/export/pdf")
async def export_to_pdf(
    content_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Export fetched content as PDF report."""
    content = db.query(FetchedContent).filter(
        FetchedContent.id == content_id,
        FetchedContent.user_id == current_user.id
    ).first()

    if not content:
        raise HTTPException(status_code=404, detail="Content not found")

    try:
        pdf_bytes = generate_pdf_report(
            title=content.title or "Untitled",
            content=content.content or "",
            executive_summary=content.executive_summary,
            technical_summary=content.technical_summary,
            iocs=content.iocs,
            metadata=content.metadata
        )

        filename = f"{content.title or 'report'}_{content_id}.pdf".replace(" ", "_")

        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except Exception as e:
        logger.error("pdf_export_failed", error=str(e))
        raise HTTPException(status_code=500, detail="PDF export failed")


@router.get("/{content_id}/export/word")
async def export_to_word(
    content_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Export fetched content as Word (.docx) report."""
    content = db.query(FetchedContent).filter(
        FetchedContent.id == content_id,
        FetchedContent.user_id == current_user.id
    ).first()

    if not content:
        raise HTTPException(status_code=404, detail="Content not found")

    try:
        docx_bytes = generate_word_report(
            title=content.title or "Untitled",
            content=content.content or "",
            executive_summary=content.executive_summary,
            technical_summary=content.technical_summary,
            iocs=content.iocs,
            metadata=content.metadata
        )

        filename = f"{content.title or 'report'}_{content_id}.docx".replace(" ", "_")

        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except Exception as e:
        logger.error("report_export_failed", error=str(e))
        raise HTTPException(status_code=500, detail="Report generation failed")


@router.post("/batch-export/pdf")
async def batch_export_pdf(
    content_ids: List[int],
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Export multiple content items as a single PDF report.
    Useful for generating weekly/monthly digest reports.
    """
    if len(content_ids) > 50:
        raise HTTPException(status_code=400, detail="Batch size cannot exceed 50 items")

    contents = db.query(FetchedContent).filter(
        FetchedContent.id.in_(content_ids),
        FetchedContent.user_id == current_user.id
    ).all()

    if not contents:
        raise HTTPException(status_code=404, detail="No content found")

    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()

        # Cover page
        story.append(Paragraph(f"Content Digest Report", styles['Title']))
        story.append(Paragraph(f"{len(contents)} Articles", styles['Heading2']))
        story.append(Paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", styles['Normal']))
        story.append(PageBreak())

        # Add each content item
        for idx, content in enumerate(contents, 1):
            story.append(Paragraph(f"{idx}. {content.title or 'Untitled'}", styles['Heading1']))

            if content.executive_summary:
                story.append(Paragraph("<b>Executive Summary</b>", styles['Heading2']))
                story.append(Paragraph(content.executive_summary, styles['BodyText']))
                story.append(Spacer(1, 0.2 * inch))

            if content.technical_summary:
                story.append(Paragraph("<b>Technical Summary</b>", styles['Heading2']))
                story.append(Paragraph(content.technical_summary, styles['BodyText']))

            if idx < len(contents):
                story.append(PageBreak())

        doc.build(story)
        buffer.seek(0)

        filename = f"digest_report_{datetime.utcnow().strftime('%Y%m%d')}.pdf"

        return Response(
            content=buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except Exception as e:
        logger.error("batch_export_failed", error=str(e))
        raise HTTPException(status_code=500, detail="Batch export failed")
