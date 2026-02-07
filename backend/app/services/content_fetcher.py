"""
Multi-format content fetcher service.
Supports fetching and parsing content from HTML, PDF, Word, CSV, and other formats.
"""
import httpx
from typing import Dict, Optional, Any
from urllib.parse import urlparse
import mimetypes
from bs4 import BeautifulSoup
from app.core.logging import logger


class ContentFetchError(Exception):
    """Raised when content fetching fails."""
    pass


class ContentFetcherService:
    """Service for fetching and parsing multi-format content."""

    def __init__(self, timeout: int = 30):
        self.timeout = timeout
        self.user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"

    async def fetch_content(self, url: str) -> Dict[str, Any]:
        """
        Fetch content from a URL and parse based on content type.

        Args:
            url: The URL to fetch

        Returns:
            Dict containing:
            - title: str
            - content: str (extracted text)
            - content_format: str (html, pdf, docx, csv, txt, etc.)
            - metadata: dict (format-specific metadata)

        Raises:
            ContentFetchError: If fetching or parsing fails
        """
        try:
            # Fetch the content
            async with httpx.AsyncClient(timeout=self.timeout, follow_redirects=True) as client:
                headers = {"User-Agent": self.user_agent}
                response = await client.get(url, headers=headers)
                response.raise_for_status()

            # Detect content type
            content_type = response.headers.get("content-type", "").lower()
            content_format = self._detect_format(url, content_type)

            logger.info(
                "content_fetched",
                url=url,
                content_type=content_type,
                detected_format=content_format,
                size_bytes=len(response.content)
            )

            # Parse based on format
            if content_format == "html":
                return await self._parse_html(response.content, url)
            elif content_format == "pdf":
                return await self._parse_pdf(response.content, url)
            elif content_format == "docx":
                return await self._parse_docx(response.content, url)
            elif content_format == "csv":
                return await self._parse_csv(response.content, url)
            elif content_format == "xlsx":
                return await self._parse_excel(response.content, url)
            elif content_format == "txt":
                return await self._parse_text(response.content, url)
            else:
                # Default to text parsing
                return await self._parse_text(response.content, url)

        except httpx.HTTPStatusError as e:
            raise ContentFetchError(f"HTTP error {e.response.status_code}: {e.response.text[:200]}")
        except httpx.TimeoutException:
            raise ContentFetchError(f"Request timeout after {self.timeout}s")
        except Exception as e:
            logger.error("content_fetch_failed", url=url, error=str(e))
            raise ContentFetchError(f"Failed to fetch content: {str(e)}")

    def _detect_format(self, url: str, content_type: str) -> str:
        """Detect content format from URL and content type."""
        # Check content type header first
        if "text/html" in content_type:
            return "html"
        elif "application/pdf" in content_type:
            return "pdf"
        elif "application/vnd.openxmlformats-officedocument.wordprocessingml" in content_type:
            return "docx"
        elif "text/csv" in content_type:
            return "csv"
        elif "application/vnd.openxmlformats-officedocument.spreadsheetml" in content_type:
            return "xlsx"
        elif "application/vnd.ms-excel" in content_type:
            return "xlsx"
        elif "text/plain" in content_type:
            return "txt"

        # Fallback to URL extension
        path = urlparse(url).path.lower()
        if path.endswith(".pdf"):
            return "pdf"
        elif path.endswith(".docx"):
            return "docx"
        elif path.endswith(".doc"):
            return "doc"  # Legacy Word format
        elif path.endswith(".csv"):
            return "csv"
        elif path.endswith((".xlsx", ".xls")):
            return "xlsx"
        elif path.endswith(".txt"):
            return "txt"
        elif path.endswith((".html", ".htm")):
            return "html"

        # Default to HTML
        return "html"

    async def _parse_html(self, content: bytes, url: str) -> Dict[str, Any]:
        """Parse HTML content."""
        try:
            soup = BeautifulSoup(content, "lxml")

            # Extract title
            title = None
            if soup.title:
                title = soup.title.string
            elif soup.find("h1"):
                title = soup.find("h1").get_text(strip=True)

            # Remove script and style elements
            for script in soup(["script", "style", "nav", "header", "footer"]):
                script.decompose()

            # Extract text content
            text_content = soup.get_text(separator="\n", strip=True)

            # Extract metadata
            metadata = {
                "url": url,
                "title_tag": soup.title.string if soup.title else None,
                "meta_description": None,
                "meta_keywords": None,
            }

            # Extract meta tags
            meta_desc = soup.find("meta", attrs={"name": "description"})
            if meta_desc:
                metadata["meta_description"] = meta_desc.get("content")

            meta_keywords = soup.find("meta", attrs={"name": "keywords"})
            if meta_keywords:
                metadata["meta_keywords"] = meta_keywords.get("content")

            return {
                "title": title or "Untitled",
                "content": text_content,
                "content_format": "html",
                "metadata": metadata
            }

        except Exception as e:
            logger.error("html_parse_failed", url=url, error=str(e))
            raise ContentFetchError(f"Failed to parse HTML: {str(e)}")

    async def _parse_pdf(self, content: bytes, url: str) -> Dict[str, Any]:
        """Parse PDF content."""
        try:
            from pypdf import PdfReader
            from io import BytesIO

            pdf_file = BytesIO(content)
            pdf_reader = PdfReader(pdf_file)

            # Extract text from all pages
            text_content = []
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())

            # Extract metadata
            metadata = {
                "url": url,
                "page_count": len(pdf_reader.pages),
                "pdf_metadata": pdf_reader.metadata if pdf_reader.metadata else {}
            }

            # Try to get title from PDF metadata
            title = "Untitled PDF"
            if pdf_reader.metadata and pdf_reader.metadata.get("/Title"):
                title = pdf_reader.metadata.get("/Title")

            return {
                "title": title,
                "content": "\n\n".join(text_content),
                "content_format": "pdf",
                "metadata": metadata
            }

        except ImportError:
            raise ContentFetchError("pypdf library not installed. Install with: pip install pypdf")
        except Exception as e:
            logger.error("pdf_parse_failed", url=url, error=str(e))
            raise ContentFetchError(f"Failed to parse PDF: {str(e)}")

    async def _parse_docx(self, content: bytes, url: str) -> Dict[str, Any]:
        """Parse Word (.docx) content."""
        try:
            from docx import Document
            from io import BytesIO

            docx_file = BytesIO(content)
            doc = Document(docx_file)

            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extract title (first heading or first paragraph)
            title = "Untitled Document"
            if paragraphs:
                title = paragraphs[0][:100]  # First 100 chars

            # Extract metadata
            metadata = {
                "url": url,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }

            # Try to get core properties
            if doc.core_properties:
                metadata["author"] = doc.core_properties.author
                metadata["title_property"] = doc.core_properties.title
                if doc.core_properties.title:
                    title = doc.core_properties.title

            return {
                "title": title,
                "content": "\n\n".join(paragraphs),
                "content_format": "docx",
                "metadata": metadata
            }

        except ImportError:
            raise ContentFetchError("python-docx library not installed. Install with: pip install python-docx")
        except Exception as e:
            logger.error("docx_parse_failed", url=url, error=str(e))
            raise ContentFetchError(f"Failed to parse Word document: {str(e)}")

    async def _parse_csv(self, content: bytes, url: str) -> Dict[str, Any]:
        """Parse CSV content."""
        try:
            import csv
            from io import StringIO

            # Decode content
            text_content = content.decode("utf-8-sig")  # Handle BOM
            csv_reader = csv.reader(StringIO(text_content))

            rows = list(csv_reader)

            # Extract title from filename or first row
            filename = urlparse(url).path.split("/")[-1]
            title = filename or "CSV Data"

            # Format as text table
            formatted_content = []
            for row in rows:
                formatted_content.append(" | ".join(row))

            metadata = {
                "url": url,
                "row_count": len(rows),
                "column_count": len(rows[0]) if rows else 0,
            }

            return {
                "title": title,
                "content": "\n".join(formatted_content),
                "content_format": "csv",
                "metadata": metadata
            }

        except Exception as e:
            logger.error("csv_parse_failed", url=url, error=str(e))
            raise ContentFetchError(f"Failed to parse CSV: {str(e)}")

    async def _parse_excel(self, content: bytes, url: str) -> Dict[str, Any]:
        """Parse Excel (.xlsx, .xls) content."""
        try:
            from openpyxl import load_workbook
            from io import BytesIO

            excel_file = BytesIO(content)
            workbook = load_workbook(excel_file, data_only=True)

            # Extract text from all sheets
            formatted_content = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                formatted_content.append(f"=== Sheet: {sheet_name} ===\n")

                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    formatted_content.append(row_text)

                formatted_content.append("")  # Blank line between sheets

            # Extract title
            filename = urlparse(url).path.split("/")[-1]
            title = filename or "Excel Data"

            metadata = {
                "url": url,
                "sheet_count": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames,
            }

            return {
                "title": title,
                "content": "\n".join(formatted_content),
                "content_format": "xlsx",
                "metadata": metadata
            }

        except ImportError:
            raise ContentFetchError("openpyxl library not installed. Install with: pip install openpyxl")
        except Exception as e:
            logger.error("excel_parse_failed", url=url, error=str(e))
            raise ContentFetchError(f"Failed to parse Excel: {str(e)}")

    async def _parse_text(self, content: bytes, url: str) -> Dict[str, Any]:
        """Parse plain text content."""
        try:
            # Try UTF-8 first, fallback to latin-1
            try:
                text_content = content.decode("utf-8")
            except UnicodeDecodeError:
                text_content = content.decode("latin-1")

            # Extract title from filename or first line
            filename = urlparse(url).path.split("/")[-1]
            first_line = text_content.split("\n")[0][:100] if text_content else ""
            title = filename or first_line or "Text Document"

            metadata = {
                "url": url,
                "character_count": len(text_content),
                "line_count": text_content.count("\n"),
            }

            return {
                "title": title,
                "content": text_content,
                "content_format": "txt",
                "metadata": metadata
            }

        except Exception as e:
            logger.error("text_parse_failed", url=url, error=str(e))
            raise ContentFetchError(f"Failed to parse text: {str(e)}")


# Singleton instance
content_fetcher = ContentFetcherService()
